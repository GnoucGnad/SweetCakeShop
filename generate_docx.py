from docx import Document
import sys

def create_docx():
    doc = Document()
    doc.add_heading('BÁO CÁO TIẾN ĐỘ ĐỒ ÁN 1 (GR1)', 0)
    
    doc.add_paragraph('Tên đề tài: Xây dựng Website Thương mại Điện tử - Cửa hàng Bánh ngọt "Sweet Cake Shop"')
    doc.add_paragraph('Sinh viên thực hiện: Đặng Đình Cường (MSSV: 20225795)')
    doc.add_paragraph('Giảng viên hướng dẫn: TS. Nguyễn Khanh Văn')
    
    doc.add_heading('1. MỤC TIÊU CỦA ĐỒ ÁN', level=1)
    doc.add_paragraph('Đồ án tập trung vào việc xây dựng một hệ thống website bán bánh ngọt hoàn chỉnh với các yêu cầu:')
    doc.add_paragraph('- Giao diện (Frontend): Thiết kế mượt mà, chuyên nghiệp, lấy cảm hứng từ các thương hiệu lớn như FreshGarden. Giao diện trực quan, dễ sử dụng, tương thích trên nhiều thiết bị (Responsive design).')
    doc.add_paragraph('- Hệ thống xử lý (Backend): Xây dựng RESTful API chuẩn mực bằng Spring Boot (Java) kết nối mượt mà với cơ sở dữ liệu.')
    doc.add_paragraph('- Cơ sở dữ liệu (Database): Thiết kế logic bằng PostgreSQL, sử dụng triggers và functions để tự động hóa nghiệp vụ (như quản lý số lượng tồn kho).')
    
    doc.add_heading('2. KẾT QUẢ HIỆN TẠI VÀ QUÁ TRÌNH THỰC HIỆN', level=1)
    
    doc.add_heading('2.1. Phân tích và Thiết kế Cơ sở Dữ liệu (Hoàn thành)', level=2)
    doc.add_paragraph('- Thiết kế mô hình ERD với 7 bảng chính: customer, customerorder, orderdetail, payment, product, optioncake, review.')
    doc.add_paragraph('- Triển khai thành công trên PostgreSQL 17.')
    doc.add_paragraph('- Xây dựng các triggers quan trọng: decrease_product_quantity() và restore_product_quantity().')
    
    doc.add_heading('2.2. Xây dựng Backend API (Hoàn thành cốt lõi)', level=2)
    doc.add_paragraph('- Xây dựng kiến trúc Spring Boot (Java 17).')
    doc.add_paragraph('- Khởi tạo toàn bộ Entity models ánh xạ trực tiếp từ CSDL.')
    doc.add_paragraph('- Xây dựng các REST Controllers (Product, Category, Order, Review) xử lý API mượt mà.')
    
    doc.add_heading('2.3. Thiết kế và Tích hợp Frontend UI (Hoàn thiện Giai đoạn 1)', level=2)
    doc.add_paragraph('Giao diện đã được thiết kế hoàn thiện, bám sát cấu trúc của các trang thương mại điện tử chuyên nghiệp (như Fresh Garden).')
    doc.add_paragraph('- Xây dựng Design System (style.css, pages.css): Quy chuẩn hóa typography, bảng màu ấm áp, hiệu ứng hover, CSS Grid/Flexbox mượt mà.')
    doc.add_paragraph('- Xây dựng 5 trang chức năng chính: Trang chủ, Sản phẩm, Chi tiết sản phẩm, Giỏ hàng, Thanh toán.')
    doc.add_paragraph('- Áp dụng JS thuần để quản lý LocalStorage (Giỏ hàng) và fetch data từ API nhanh chóng.')
    
    doc.add_heading('2.4. Kỹ năng Định hướng và Quản trị Trợ lý Trí tuệ Nhân tạo (AI Prompt Engineering)', level=2)
    doc.add_paragraph('Trong bối cảnh công nghệ hiện đại, vai trò của một kỹ sư phần mềm không chỉ dừng lại ở việc trực tiếp viết mã lệnh (hard-coding), mà còn vươn tầm lên việc thiết kế kiến trúc hệ thống, đánh giá rủi ro và điều phối, lãnh đạo các công cụ AI. Đồ án này là minh chứng thực tiễn cho khả năng ứng dụng AI như một "cộng sự lập trình" đắc lực dưới sự kiểm soát chặt chẽ của người kỹ sư.')
    
    doc.add_paragraph('- Thiết lập quy tắc và định hướng kiến trúc (Guidelines & Architecture Setup):')
    doc.add_paragraph('  + Ngay từ giai đoạn khởi tạo, tôi đã đóng vai trò Technical Lead, đưa ra các "ràng buộc" nghiêm ngặt mang tính nền tảng.')
    doc.add_paragraph('  + Yêu cầu AI bắt buộc triển khai bằng hệ quản trị PostgreSQL thay vì các cơ sở dữ liệu phi quan hệ, buộc AI phải thiết kế frontend bằng CSS thuần (Vanilla CSS) thay vì lạm dụng các framework có sẵn (như Tailwind hay Bootstrap) nhằm rèn luyện kỹ năng kiểm soát DOM và tối ưu hóa hiệu năng tải trang.')
    doc.add_paragraph('  + AI phải tuân thủ tuyệt đối quy trình thác nước lai Agile: Hoàn thiện Database -> Đóng gói Backend API -> Triển khai Frontend.')
    
    doc.add_paragraph('- Sự cố, rào cản từ AI và Biện pháp khắc phục (AI Hallucinations & Code Review):')
    doc.add_paragraph('  + Sự cố mất kiểm soát tính đồng nhất (Inconsistency): Trong quá trình chia nhỏ các module Frontend, AI có xu hướng "lười biếng" bằng cách viết CSS trực tiếp (inline-style) hoặc can thiệp sai file, phá vỡ cấu trúc Grid/Flexbox sẵn có (ví dụ: lỗi tràn ô nhập số lượng ở Giỏ hàng).')
    doc.add_paragraph('    > Khắc phục: Tôi phải lập tức can thiệp bằng các "Prompt" mang tính mệnh lệnh cứng rắn. Đóng vai trò là một Code Reviewer khắt khe, tôi yêu cầu AI phải giải trình logic trước khi sửa, buộc AI rã cấu trúc thành các file CSS độc lập (style.css, pages.css) và chỉ định đích danh class cần thao tác để đảm bảo tính module hóa.')
    doc.add_paragraph('  + Sự cố rập khuôn và thiếu linh hoạt xử lý tệp tin: Khi chèn các ảnh sản phẩm thực tế có phông nền phức tạp, AI đưa ra các phương án CSS hòa trộn màu (mix-blend-mode) rập khuôn và hoàn toàn thất bại với ảnh nền tối.')
    doc.add_paragraph('    > Khắc phục: Tôi trực tiếp tái định hướng tư duy cho AI, yêu cầu AI dừng việc can thiệp CSS và chuyển hướng sang giải pháp Backend/Scripting. Tôi yêu cầu AI phải viết các đoạn mã Python tích hợp mạng nơ-ron (Mô hình U-2-Net) để tự động hóa việc bóc tách phông nền và tối ưu hóa nén ảnh (Image Optimization). Chuyển AI từ việc "chữa cháy UI" sang "giải quyết tận gốc bằng thuật toán".')
    
    doc.add_paragraph('- Kinh nghiệm rút ra:')
    doc.add_paragraph('  + Để "thuần hóa" và khai thác tối đa tiềm năng của AI, kỹ sư phải có tư duy hệ thống vững vàng.')
    doc.add_paragraph('  + Nhận thức rõ ràng rằng AI chỉ là một công cụ thực thi xuất sắc, nhưng "linh hồn", tính đúng đắn, và tư duy thiết kế của một sản phẩm phần mềm bắt buộc phải xuất phát từ khối óc định hướng và nhãn quan bao quát của người kỹ sư.')

    doc.add_heading('3. CÁC KHÓ KHĂN KỸ THUẬT VÀ GIẢI PHÁP', level=1)
    
    doc.add_paragraph('Ngoài việc định hướng AI, tôi cũng phải trực tiếp giám sát và đưa ra giải pháp cho các vấn đề kỹ thuật chuyên sâu:')
    
    doc.add_paragraph('- Khó khăn 1: Xử lý phông nền ảnh (Background Removal)')
    doc.add_paragraph('  + Vấn đề: Các ảnh tĩnh tải lên (Slide, Bánh) có phông nền trắng hoặc xám/đen, khi chèn vào website làm vỡ bố cục màu sắc.')
    doc.add_paragraph('  + Giải pháp: Xây dựng script tự động bóc tách phông nền bằng AI U-2-Net (thư viện rembg), xuất ra PNG trong suốt giúp giao diện liền mạch.')
    
    doc.add_paragraph('- Khó khăn 2: Tối ưu hóa hiệu năng (Web Performance)')
    doc.add_paragraph('  + Vấn đề: Sau khi tải lượng lớn ảnh gốc (lên tới 20MB/ảnh), website bị sụt giảm FPS và lag khi cuộn trang.')
    doc.add_paragraph('  + Giải pháp: Viết script Python (Pillow) để tự động Resize (giới hạn 1920px) và Compress (thuật toán Lanczos, chất lượng 85%), giảm dung lượng file xuống gấp 10 lần (từ 18MB xuống 1.9MB), đưa tốc độ tải trang về mức hoàn hảo.')
    
    doc.add_heading('4. MỤC TIÊU TIẾP THEO (ĐỊNH HƯỚNG TƯƠNG LAI)', level=1)
    doc.add_paragraph('Mặc dù Frontend đã hoàn thiện, hệ thống vẫn cần phát triển thêm để trở thành một sản phẩm có thể "sống" và vận hành thực tế:')
    doc.add_paragraph('- Tích hợp Thanh toán trực tuyến: Kết nối với các cổng thanh toán điện tử thực tế như VNPay, MoMo, hoặc PayPal.')
    doc.add_paragraph('- Hệ thống Xác thực (Authentication): Xây dựng tính năng Đăng nhập/Đăng ký cho Khách hàng và Admin bằng Spring Security & JWT.')
    doc.add_paragraph('- Xây dựng Admin Dashboard: Trang quản trị nội bộ dành cho chủ cửa hàng để Quản lý và Thống kê doanh thu.')
    doc.add_paragraph('- Deploy (Triển khai): Đưa hệ thống lên môi trường đám mây để có thể truy cập qua domain công khai.')
    
    doc.save(r'd:\GR1\baocao_tiendo_gr1_v3.docx')
    print("Docx created successfully.")

if __name__ == "__main__":
    create_docx()
